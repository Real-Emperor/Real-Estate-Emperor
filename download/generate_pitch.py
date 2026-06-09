from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ─── Page setup ───
for section in doc.sections:
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

# ─── Color palette ───
GREEN_DARK  = RGBColor(0x1B, 0x5E, 0x20)
GREEN_MID   = RGBColor(0x2E, 0x7D, 0x32)
GREEN_LIGHT = RGBColor(0x4C, 0xAF, 0x50)
GOLD        = RGBColor(0xC6, 0x8A, 0x1D)
GOLD_DARK   = RGBColor(0xA0, 0x6F, 0x10)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
BLACK       = RGBColor(0x00, 0x00, 0x00)
GRAY        = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY  = RGBColor(0x99, 0x99, 0x99)
RED_STRIKE  = RGBColor(0xC6, 0x28, 0x28)

# ─── Helper functions ───
def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_styled_paragraph(doc_or_cell, text, font_size=11, bold=False, color=BLACK, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0), font_name='Calibri'):
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color
    run.font.name = font_name
    return p

def add_mixed_paragraph(doc_or_cell, parts, alignment=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6), space_before=Pt(0)):
    """parts = list of (text, font_size, bold, color, font_name)"""
    p = doc_or_cell.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    for text, font_size, bold, color, font_name in parts:
        run = p.add_run(text)
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.color.rgb = color
        run.font.name = font_name
    return p

def add_gold_line(doc):
    """Add a thin gold horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    # Use a border-bottom on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="C68A1D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_green_line(doc):
    """Add a thin green horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1B5E20"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p

def add_section_header(doc, text, color=GREEN_DARK, font_name='Calibri'):
    """Add a big section header with underline."""
    add_green_line(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = color
    run.font.name = font_name
    add_gold_line(doc)
    return p

def add_bullet(doc, text, font_size=11, color=BLACK, font_name='Calibri'):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    # Clear default run and add our own
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.color.rgb = color
    run.font.name = font_name
    return p


# ═══════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════

# Top green bar
cover_table = doc.add_table(rows=1, cols=1)
cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cover_table.cell(0, 0)
set_cell_shading(cell, "1B5E20")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
p.paragraph_format.space_after = Pt(30)
run = p.add_run("REAL ESTATE EMPEROR")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'
p2 = cell.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(4)
run2 = p2.add_run("Real Estate & General Maintenance L.L.C.")
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0xC6, 0x8A, 0x1D)
run2.font.name = 'Calibri'
p3 = cell.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_after = Pt(20)
run3 = p3.add_run("الإمبراطور العقاري لإدارة الممتلكات ذ.م.م")
run3.font.size = Pt(14)
run3.font.color.rgb = WHITE
run3.font.name = 'Calibri'

# Spacer
doc.add_paragraph().paragraph_format.space_after = Pt(10)

# Middle - Presentation title
add_styled_paragraph(doc, "BUSINESS PROPOSAL", font_size=22, bold=True, color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_styled_paragraph(doc, "ব্যবসায়িক প্রস্তাবনা", font_size=18, bold=True, color=GOLD_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))

add_gold_line(doc)

add_styled_paragraph(doc, "Smart Property Management Dashboard", font_size=14, bold=False, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_styled_paragraph(doc, "স্মার্ট সম্পত্তি ব্যবস্থাপনা ড্যাশবোর্ড", font_size=13, bold=False, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

# QR Code
if os.path.exists('/home/z/my-project/download/qr_code.png'):
    p_qr = doc.add_paragraph()
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr.paragraph_format.space_after = Pt(4)
    run_qr = p_qr.add_run()
    run_qr.add_picture('/home/z/my-project/download/qr_code.png', width=Inches(1.8))

    add_styled_paragraph(doc, "Scan to view the live dashboard", font_size=9, bold=False, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
    add_styled_paragraph(doc, "real-estate-emperor.vercel.app", font_size=9, bold=True, color=GREEN_MID, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

# Prepared for
add_green_line(doc)
add_styled_paragraph(doc, "Prepared for:", font_size=10, bold=False, color=GRAY, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2), space_before=Pt(10))
add_styled_paragraph(doc, "Shafiul Azam Alhaj Abdul Sukkur", font_size=16, bold=True, color=GREEN_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
add_styled_paragraph(doc, "শফিউল আজম আলহাজ আব্দুল সুক্কুর", font_size=14, bold=True, color=GOLD_DARK, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))

# Page break after cover
doc.add_page_break()


# ═══════════════════════════════════════════════════════════════
# ENGLISH SECTION
# ═══════════════════════════════════════════════════════════════

# English header bar
en_table = doc.add_table(rows=1, cols=1)
en_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = en_table.cell(0, 0)
set_cell_shading(cell, "1B5E20")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(14)
run = p.add_run("ENGLISH  ·  BUSINESS PROPOSAL")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Greeting
add_styled_paragraph(doc, "Dear Shafiul Azam Alhaj Abdul Sukkur,", font_size=12, bold=False, color=BLACK, space_after=Pt(10))

add_styled_paragraph(doc, 
    "You built this company from nothing in a foreign land. Every dirham in your portfolio was earned through sacrifice, discipline, and trust in Allah. Now you manage dozens of tenants across multiple properties — and the rent tracking is still in notebooks, WhatsApp groups, and your memory. That works until it doesn't. One missed payment, one forgotten follow-up, one miscommunication — and you lose money that was rightfully yours.",
    font_size=11, color=BLACK, space_after=Pt(10))

add_styled_paragraph(doc,
    "I built something for you. Not a generic app — a dashboard tailored to Real Estate Emperor, with your company's identity, your languages, and your way of doing business baked into every screen. It's already live. It's already working.",
    font_size=11, color=BLACK, space_after=Pt(14))

# ─── The Problem ───
add_section_header(doc, "THE PROBLEM")

add_styled_paragraph(doc,
    "Right now, tracking who paid and who didn't means scrolling through WhatsApp chats, flipping through paper, and relying on memory. When a tenant is late, you or your staff have to manually compose messages, remember what was said last month, and hope nothing falls through the cracks. Multiply that across dozens of units and you have a system that guarantees lost revenue — not because tenants won't pay, but because the process makes it too easy for things to slip.",
    font_size=11, color=BLACK, space_after=Pt(10))

# ─── The Solution ───
add_section_header(doc, "THE SOLUTION")

add_styled_paragraph(doc,
    "The Real Estate Emperor Dashboard puts your entire rental operation on one screen. Every property, every tenant, every payment — organized, color-coded, and actionable. Green means paid. Red means overdue. One click sends a WhatsApp reminder. Your staff sees only what they need; you see everything. It's designed for the way you already work — just faster, cleaner, and impossible to forget.",
    font_size=11, color=BLACK, space_after=Pt(10))

# ─── Key Features ───
add_section_header(doc, "KEY FEATURES")

features = [
    ("Multi-Language Interface", "English, Arabic, Bengali, and Urdu — your tenants and staff use it in their own language."),
    ("Tenant Scoring System", "Automatically ranks tenants by payment reliability so you know who to trust and who to watch."),
    ("WhatsApp One-Click Reminders", "Late payment? One tap sends a professional reminder directly to their WhatsApp — no typing, no excuses."),
    ("Property Lifecycle Tracking", "From vacant → occupied → maintenance → renewed — see exactly where every unit stands."),
    ("Red/Green Payment Status", "Visual system: green = paid, red = overdue. You see problems instantly without reading a single number."),
    ("Role-Based Access", "Your staff sees only their tasks. You see everything. Full control, zero confusion."),
    ("Mobile-First Design", "Built for your phone first. Check your portfolio standing in line at the grocery store."),
    ("Islamic Bengali Identity", "Your brand, your culture, your values — not a generic white-label template."),
]

for title, desc in features:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run_t = p.add_run(f"► {title}: ")
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = GREEN_DARK
    run_t.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = BLACK
    run_d.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── Pricing ───
add_section_header(doc, "INVESTMENT")

# Pricing table
pricing_table = doc.add_table(rows=4, cols=3)
pricing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
pricing_table.style = 'Table Grid'

# Set column widths
for row in pricing_table.rows:
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(1.5)

# Header row
headers = ["", "Standard Rate", "Your Rate"]
for i, h in enumerate(headers):
    cell = pricing_table.cell(0, i)
    set_cell_shading(cell, "1B5E20")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'

# Setup fee row
setup_labels = ["One-Time Setup", "5,000 AED", "3,000 AED"]
for i, label in enumerate(setup_labels):
    cell = pricing_table.cell(1, i)
    if i == 0:
        set_cell_shading(cell, "F1F8E9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    if i == 1:
        run.font.color.rgb = RED_STRIKE
        run.font.strike = True
    elif i == 2:
        run.font.bold = True
        run.font.color.rgb = GREEN_DARK
        run.font.size = Pt(14)

# Monthly fee row
monthly_labels = ["Monthly (Hosting, Maintenance & Updates)", "1,500 AED", "1,000 AED"]
for i, label in enumerate(monthly_labels):
    cell = pricing_table.cell(2, i)
    if i == 0:
        set_cell_shading(cell, "F1F8E9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    if i == 1:
        run.font.color.rgb = RED_STRIKE
        run.font.strike = True
    elif i == 2:
        run.font.bold = True
        run.font.color.rgb = GREEN_DARK
        run.font.size = Pt(14)

# Savings row
cell = pricing_table.cell(3, 0)
set_cell_shading(cell, "C68A1D")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("YOUR SAVINGS")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

# Merge cells 1 and 2 for savings value
cell_savings = pricing_table.cell(3, 1)
cell_savings_r = pricing_table.cell(3, 2)
cell_savings.merge(cell_savings_r)
set_cell_shading(cell_savings, "FFF8E1")
p = cell_savings.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("2,000 AED setup + 500 AED/month = 8,000 AED saved in Year 1")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = GOLD_DARK
run.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── No Lock-In ───
add_section_header(doc, "NO LOCK-IN. NO RISK.")

add_styled_paragraph(doc,
    "There is no long-term contract. No legal obligation on either side. You can terminate at any time — just stop paying and the service stops. I believe in this product enough to let the results speak for themselves. If it doesn't make your life easier, you walk away. That simple.",
    font_size=11, color=BLACK, space_after=Pt(12))

# ─── Call to Action ───
cta_table = doc.add_table(rows=1, cols=1)
cta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cta_table.cell(0, 0)
set_cell_shading(cell, "1B5E20")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(16)
run = p.add_run("Let's talk. The dashboard is already live.\nScan the QR code or visit: real-estate-emperor.vercel.app")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

# QR code at bottom of English section
if os.path.exists('/home/z/my-project/download/qr_code.png'):
    p_qr2 = doc.add_paragraph()
    p_qr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr2.paragraph_format.space_before = Pt(10)
    p_qr2.paragraph_format.space_after = Pt(2)
    run_qr2 = p_qr2.add_run()
    run_qr2.add_picture('/home/z/my-project/download/qr_code.png', width=Inches(1.3))

# ═══════════════════════════════════════════════════════════════
# BENGALI SECTION — EQUAL DESIGN
# ═══════════════════════════════════════════════════════════════
doc.add_page_break()

# Bengali header bar
bn_table = doc.add_table(rows=1, cols=1)
bn_table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = bn_table.cell(0, 0)
set_cell_shading(cell, "1B5E20")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14)
p.paragraph_format.space_after = Pt(14)
run = p.add_run("বাংলা  ·  ব্যবসায়িক প্রস্তাবনা")
run.font.size = Pt(16)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# Bengali Greeting
add_styled_paragraph(doc, "শ্রদ্ধেয় শফিউল আজম আলহাজ আব্দুল সুক্কুর সাহেব,", font_size=12, bold=False, color=BLACK, space_after=Pt(10))

add_styled_paragraph(doc,
    "আপনি এই কোম্পানি বিদেশের মাটিতে শূন্য থেকে গড়ে তুলেছেন। আপনার পোর্টফোলিওর প্রতিটি দিরহাম অর্জিত হয়েছে আত্মত্যাগ, সংযম এবং আল্লাহর উপর ভরসার মাধ্যমে। এখন আপনি একাধিক সম্পত্তিতে কয়েক ডজন ভাড়াটিয়া পরিচালনা করেন — আর ভাড়া ট্র্যাকিং এখনও খাতায়, হোয়াটসঅ্যাপ গ্রুপে এবং আপনার স্মৃতিতে। এটা কাজ করে — যতক্ষণ না কাজ করে। একটি মিসড পেমেন্ট, একটি ভুলে যাওয়া ফলো-আপ, একটি ভুল যোগাযোগ — এবং আপনি এমন টাকা হারালেন যা আপনার অধিকারে ছিল।",
    font_size=11, color=BLACK, space_after=Pt(10))

add_styled_paragraph(doc,
    "আমি আপনার জন্য কিছু তৈরি করেছি। কোনো সাধারণ অ্যাপ নয় — আল রীফ আল জুনুবির জন্য সম্পূর্ণ কাস্টমাইজড একটি ড্যাশবোর্ড, যেখানে আপনার কোম্পানির পরিচয়, আপনার ভাষা এবং আপনার ব্যবসায়িক পদ্ধতি প্রতিটি স্ক্রিনে অন্তর্ভুক্ত। এটি ইতিমধ্যে লাইভ। ইতিমধ্যে কাজ করছে।",
    font_size=11, color=BLACK, space_after=Pt(14))

# ─── সমস্যা ───
add_section_header(doc, "সমস্যা")

add_styled_paragraph(doc,
    "এখন কে টাকা দিয়েছে আর কে দেয়নি তা ট্র্যাক করার মানে হলো হোয়াটসঅ্যাপ চ্যাট স্ক্রল করা, কাগজ ঘাঁটা এবং স্মৃতির উপর নির্ভর করা। কোনো ভাড়াটিয়া দেরি করলে আপনাকে বা আপনার কর্মীকে ম্যানুয়ালি মেসেজ লিখতে হয়, গত মাসে কী বলা হয়েছিল মনে রাখতে হয়, এবং আশা করতে হয় কিছু বাদ পড়বে না। ডজনখানেক ইউনিটে এটি গুণ করলে আপনি এমন একটি সিস্টেম পান যা গ্যারান্টি দেয় রাজস্ব হারাবেন — ভাড়াটিয়ারা দেবে না বলে নয়, বরং প্রক্রিয়াটি এতটাই জটিল যে জিনিসগুলো ফসকে যাওয়া সহজ।",
    font_size=11, color=BLACK, space_after=Pt(10))

# ─── সমাধান ───
add_section_header(doc, "সমাধান")

add_styled_paragraph(doc,
    "আল রীফ আল জুনুবি ড্যাশবোর্ড আপনার সম্পূর্ণ ভাড়া ব্যবস্থাপনাকে একটি স্ক্রিনে নিয়ে আসে। প্রতিটি সম্পত্তি, প্রতিটি ভাড়াটিয়া, প্রতিটি পেমেন্ট — সুসংগঠিত, রঙ-কোডেড এবং কার্যকর। সবুজ মানে পরিশোধিত। লাল মানে বকেয়া। এক ক্লিকে হোয়াটসঅ্যাপ রিমাইন্ডার পাঠান। আপনার কর্মীরা শুধু তাদের প্রয়োজনীয় জিনিস দেখেন; আপনি সব দেখেন। এটি আপনার বর্তমান কাজের পদ্ধতির জন্যই ডিজাইন করা — শুধু আরও দ্রুত, পরিষ্কার এবং ভুলার অসম্ভব।",
    font_size=11, color=BLACK, space_after=Pt(10))

# ─── মূল বৈশিষ্ট্য ───
add_section_header(doc, "মূল বৈশিষ্ট্য")

bn_features = [
    ("বহুভাষিক ইন্টারফেস", "ইংরেজি, আরবি, বাংলা এবং উর্দু — আপনার ভাড়াটিয়া এবং কর্মীরা নিজ নিজ ভাষায় ব্যবহার করতে পারবেন।"),
    ("ভাড়াটিয়া স্কোরিং সিস্টেম", "পেমেন্ট নির্ভরযোগ্যতা অনুযায়ী স্বয়ংক্রিয়ভাবে ভাড়াটিয়াদের র‍্যাংক করে যাতে আপনি জানতে পারেন কাকে বিশ্বাস করবেন আর কাকে নজরে রাখবেন।"),
    ("হোয়াটসঅ্যাপ ওয়ান-ক্লিক রিমাইন্ডার", "দেরিতে পেমেন্ট? এক ট্যাপে সরাসরি তাদের হোয়াটসঅ্যাপে পেশাদার রিমাইন্ডার পাঠান — কোনো টাইপিং নয়, কোনো অজুহাত নয়।"),
    ("সম্পত্তি লাইফসাইকেল ট্র্যাকিং", "খালি → দখলকৃত → রক্ষণাবেক্ষণ → নবায়ন — প্রতিটি ইউনিটের অবস্থা স্পষ্টভাবে দেখুন।"),
    ("লাল/সবুজ পেমেন্ট স্ট্যাটাস", "দৃশ্যমান সিস্টেম: সবুজ = পরিশোধিত, লাল = বকেয়া। একটি সংখ্যাও না পড়ে সমস্যা তৎক্ষণাৎ দেখুন।"),
    ("ভূমিকা-ভিত্তিক অ্যাক্সেস", "আপনার কর্মীরা শুধু তাদের কাজ দেখেন। আপনি সবকিছু দেখেন। সম্পূর্ণ নিয়ন্ত্রণ, শূন্য বিভ্রান্তি।"),
    ("মোবাইল-ফার্স্ট ডিজাইন", "আপনার ফোনের জন্য প্রথমে তৈরি। মুদি দোকানের লাইনে দাঁড়িয়ে আপনার পোর্টফোলিও চেক করুন।"),
    ("ইসলামিক বাঙালি পরিচয়", "আপনার ব্র্যান্ড, আপনার সংস্কৃতি, আপনার মূল্যবোধ — কোনো সাধারণ হোয়াইট-লেবেল টেমপ্লেট নয়।"),
]

for title, desc in bn_features:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(2)
    run_t = p.add_run(f"► {title}: ")
    run_t.font.size = Pt(11)
    run_t.font.bold = True
    run_t.font.color.rgb = GREEN_DARK
    run_t.font.name = 'Calibri'
    run_d = p.add_run(desc)
    run_d.font.size = Pt(10)
    run_d.font.color.rgb = BLACK
    run_d.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── মূল্য ───
add_section_header(doc, "বিনিয়োগ")

# Bengali pricing table — SAME DESIGN
pricing_table2 = doc.add_table(rows=4, cols=3)
pricing_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
pricing_table2.style = 'Table Grid'

for row in pricing_table2.rows:
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(1.5)
    row.cells[2].width = Inches(1.5)

# Header row
bn_headers = ["", "সাধারণ মূল্য", "আপনার মূল্য"]
for i, h in enumerate(bn_headers):
    cell = pricing_table2.cell(0, i)
    set_cell_shading(cell, "1B5E20")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(h)
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = WHITE
    run.font.name = 'Calibri'

# Setup fee row
bn_setup = ["এককালীন সেটআপ", "৫,০০০ AED", "৩,০০০ AED"]
for i, label in enumerate(bn_setup):
    cell = pricing_table2.cell(1, i)
    if i == 0:
        set_cell_shading(cell, "F1F8E9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    if i == 1:
        run.font.color.rgb = RED_STRIKE
        run.font.strike = True
    elif i == 2:
        run.font.bold = True
        run.font.color.rgb = GREEN_DARK
        run.font.size = Pt(14)

# Monthly fee row
bn_monthly = ["মাসিক (হোস্টিং, রক্ষণাবেক্ষণ ও আপডেট)", "১,৫০০ AED", "১,০০০ AED"]
for i, label in enumerate(bn_monthly):
    cell = pricing_table2.cell(2, i)
    if i == 0:
        set_cell_shading(cell, "F1F8E9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(label)
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    if i == 1:
        run.font.color.rgb = RED_STRIKE
        run.font.strike = True
    elif i == 2:
        run.font.bold = True
        run.font.color.rgb = GREEN_DARK
        run.font.size = Pt(14)

# Savings row
cell = pricing_table2.cell(3, 0)
set_cell_shading(cell, "C68A1D")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("আপনার সাশ্রয়")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

cell_savings2 = pricing_table2.cell(3, 1)
cell_savings2_r = pricing_table2.cell(3, 2)
cell_savings2.merge(cell_savings2_r)
set_cell_shading(cell_savings2, "FFF8E1")
p = cell_savings2.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(6)
run = p.add_run("২,০০০ AED সেটআপ + ৫০০ AED/মাস = প্রথম বছরে ৮,০০০ AED সাশ্রয়")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = GOLD_DARK
run.font.name = 'Calibri'

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── কোনো বাধ্যবাধকতা নেই ───
add_section_header(doc, "কোনো বাধ্যবাধকতা নেই। কোনো ঝুঁকি নেই।")

add_styled_paragraph(doc,
    "এখানে কোনো দীর্ঘমেয়াদী চুক্তি নেই। কোনো পক্ষের জন্যই কোনো আইনি বাধ্যবাধকতা নেই। আপনি যেকোনো সময় বাতিল করতে পারেন — শুধু পেমেন্ট বন্ধ করুন, সার্ভিস বন্ধ হয়ে যাবে। আমি এই পণ্যে যথেষ্ট বিশ্বাসী যে ফলাফল নিজেই কথা বলবে। যদি এটি আপনার জীবন সহজ না করে, আপনি চলে যান। এতটাই সহজ।",
    font_size=11, color=BLACK, space_after=Pt(12))

# ─── Bengali Call to Action ───
cta_table2 = doc.add_table(rows=1, cols=1)
cta_table2.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = cta_table2.cell(0, 0)
set_cell_shading(cell, "1B5E20")
p = cell.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
p.paragraph_format.space_after = Pt(16)
run = p.add_run("আসুন কথা বলি। ড্যাশবোর্ড ইতিমধ্যে লাইভ আছে।\nQR কোড স্ক্যান করুন অথবা ভিজিট করুন: real-estate-emperor.vercel.app")
run.font.size = Pt(13)
run.font.bold = True
run.font.color.rgb = WHITE
run.font.name = 'Calibri'

# QR code at bottom of Bengali section
if os.path.exists('/home/z/my-project/download/qr_code.png'):
    p_qr3 = doc.add_paragraph()
    p_qr3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr3.paragraph_format.space_before = Pt(10)
    p_qr3.paragraph_format.space_after = Pt(2)
    run_qr3 = p_qr3.add_run()
    run_qr3.add_picture('/home/z/my-project/download/qr_code.png', width=Inches(1.3))


# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
output_path = '/home/z/my-project/download/Real_Estate_Emperor_Pitch.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
